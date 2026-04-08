"""Export TimetableResult to a landscape A4 Word document using python-docx."""
from __future__ import annotations

import logging
import os
from datetime import datetime, timedelta
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

if TYPE_CHECKING:
    from timease.engine.models import Assignment, SchoolData, TimetableResult

logger = logging.getLogger(__name__)

_FMT = "%H:%M"
_ACCENT    = "0D9488"   # teal-600
_ACCENT_LT = "E6F4F3"   # teal-50 for alternating rows
_WHITE     = "FFFFFF"
_TEXT_DARK = "111827"   # gray-900
_ROW_ALT   = "F7F8FA"
_SEP_ROW   = "F0F0F0"
_FONT_HEAD = "Georgia"
_FONT_BODY = "Arial"


# ── XML helpers ────────────────────────────────────────────────────────────────

def _shade_cell(cell: object, hex_color: str) -> None:
    """Apply background shading to a Word table cell."""
    tc = cell._tc  # type: ignore[attr-defined]
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.strip("#").upper())
    tcPr.append(shd)


def _write_cell(
    cell: object,
    lines: list[str],
    bold: bool = False,
    font_size: int = 8,
    font_color: str = _TEXT_DARK,
    font_name: str = _FONT_BODY,
    align: str = "center",
) -> None:
    """Write multiline content to a Word table cell, replacing existing content."""
    tc = cell._tc  # type: ignore[attr-defined]
    paras = tc.findall(qn("w:p"))
    for p_el in paras[1:]:
        tc.remove(p_el)
    first_p = paras[0] if paras else OxmlElement("w:p")
    if not paras:
        tc.append(first_p)
    for r_el in first_p.findall(qn("w:r")):
        first_p.remove(r_el)

    pPr = first_p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        first_p.insert(0, pPr)
    for jc_el in pPr.findall(qn("w:jc")):
        pPr.remove(jc_el)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), align)
    pPr.append(jc)
    # Tight paragraph spacing
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "20")
    spacing.set(qn("w:after"), "20")

    for i, line in enumerate(lines):
        if i > 0:
            br_r = OxmlElement("w:r")
            br_el = OxmlElement("w:br")
            br_r.append(br_el)
            first_p.append(br_r)
        r_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        # Font
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rPr.append(rFonts)
        # Size
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(font_size * 2))
        rPr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(font_size * 2))
        rPr.append(sz_cs)
        # Bold
        if bold:
            rPr.append(OxmlElement("w:b"))
        # Color
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), font_color.strip("#").upper())
        rPr.append(color_el)
        r_el.append(rPr)
        t_el = OxmlElement("w:t")
        t_el.text = line
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_el.append(t_el)
        first_p.append(r_el)


def _make_page_field(tag: str) -> OxmlElement:
    """Build a w:fldChar / w:instrText run for PAGE or NUMPAGES."""
    run = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run.append(fld_begin)

    run2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" {tag} "
    run2.append(instr)

    run3 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3.append(fld_end)

    return run, run2, run3


def _set_header_footer(
    section,
    school_name: str,
    doc_title: str,
    gen_date: str,
) -> None:
    """Attach header (school + title) and footer (page/date/brand) to *section*."""
    section.different_first_page_header_footer = False

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    for para in list(header.paragraphs):
        p_el = para._p
        p_el.getparent().remove(p_el)

    hdr_para = header.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hdr_para.add_run(f"{school_name}  ·  {doc_title}")
    run.font.name = _FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # gray-500

    # Thin bottom border on header paragraph
    pPr = hdr_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _ACCENT)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in list(footer.paragraphs):
        p_el = para._p
        p_el.getparent().remove(p_el)

    ftr_para = footer.add_paragraph()
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Généré par TIMEASE  ·  date  ·  Page X / Y"
    run_brand = ftr_para.add_run("Généré par TIMEASE  ·  ")
    run_brand.font.name = _FONT_BODY
    run_brand.font.size = Pt(8)
    run_brand.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    run_date = ftr_para.add_run(f"{gen_date}  ·  Page ")
    run_date.font.name = _FONT_BODY
    run_date.font.size = Pt(8)
    run_date.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    r1, r2, r3 = _make_page_field("PAGE")
    ftr_para._p.append(r1)
    ftr_para._p.append(r2)
    ftr_para._p.append(r3)

    run_sep = ftr_para.add_run(" / ")
    run_sep.font.name = _FONT_BODY
    run_sep.font.size = Pt(8)
    run_sep.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    r4, r5, r6 = _make_page_field("NUMPAGES")
    ftr_para._p.append(r4)
    ftr_para._p.append(r5)
    ftr_para._p.append(r6)


# ── Cover page ────────────────────────────────────────────────────────────────

def _add_cover_page(
    doc: Document,
    data: SchoolData,
    gen_date: str,
    logo_path: str | None,
) -> None:
    """Insert a portrait cover page as the first section."""
    section = doc.sections[0]
    # Portrait A4 cover
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    # No header/footer on cover
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for hf in (section.header, section.footer):
        for para in list(hf.paragraphs):
            hf._element.remove(para._p)

    # ── Top accent band (colored paragraph) ──
    band_para = doc.add_paragraph()
    band_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = band_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "48")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _ACCENT)
        pBdr.append(el)
    pPr.append(pBdr)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), _ACCENT)
    pPr.append(shading)
    run_band = band_para.add_run("  EMPLOI DU TEMPS  ")
    run_band.font.name = _FONT_BODY
    run_band.font.size = Pt(11)
    run_band.font.bold = True
    run_band.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Vertical spacing ──
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)

    # ── Logo ──
    if logo_path and os.path.isfile(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Mm(35))

    # ── School name ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(24)
    name_run = name_para.add_run(data.school.name.upper())
    name_run.font.name = _FONT_HEAD
    name_run.font.size = Pt(28)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)  # _ACCENT

    # ── Academic year ──
    yr_para = doc.add_paragraph()
    yr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yr_para.paragraph_format.space_before = Pt(6)
    yr_run = yr_para.add_run(
        f"Année académique {data.school.academic_year}"
        if getattr(data.school, "academic_year", None)
        else "Emploi du temps scolaire"
    )
    yr_run.font.name = _FONT_BODY
    yr_run.font.size = Pt(14)
    yr_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)  # gray-600

    # ── Divider ──
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # ── Meta block ──
    city_val = getattr(data.school, "city", None) or ""
    meta_lines = [
        ("Ville", city_val),
        ("Date de génération", gen_date),
        ("Sessions planifiées", ""),
    ]
    for label, value in meta_lines:
        if not value and label == "Sessions planifiées":
            continue
        if not value:
            continue
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl_run = meta_para.add_run(f"{label} : ")
        lbl_run.font.name = _FONT_BODY
        lbl_run.font.size = Pt(11)
        lbl_run.font.bold = True
        lbl_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        val_run = meta_para.add_run(value)
        val_run.font.name = _FONT_BODY
        val_run.font.size = Pt(11)
        val_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # ── Bottom accent band ──
    for _ in range(4):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    bot_para = doc.add_paragraph()
    bot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bot_pPr = bot_para._p.get_or_add_pPr()
    bot_shd = OxmlElement("w:shd")
    bot_shd.set(qn("w:val"), "clear")
    bot_shd.set(qn("w:color"), "auto")
    bot_shd.set(qn("w:fill"), _ACCENT_LT)
    bot_pPr.append(bot_shd)
    bot_run = bot_para.add_run("timease.app")
    bot_run.font.name = _FONT_BODY
    bot_run.font.size = Pt(9)
    bot_run.font.italic = True
    bot_run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    doc.add_page_break()


# ── Timetable grid ─────────────────────────────────────────────────────────────

def _time_slots(start: str, end: str, base_min: int) -> list[tuple[str, str]]:
    """All (start, end) base-unit intervals within [start, end)."""
    cur = datetime.strptime(start, _FMT)
    stop = datetime.strptime(end, _FMT)
    delta = timedelta(minutes=base_min)
    slots: list[tuple[str, str]] = []
    while cur < stop:
        nxt = cur + delta
        slots.append((cur.strftime(_FMT), nxt.strftime(_FMT)))
        cur = nxt
    return slots


def _slot_span(start: str, end: str, base_min: int) -> int:
    """Number of base-unit rows an assignment occupies."""
    s = datetime.strptime(start, _FMT)
    e = datetime.strptime(end, _FMT)
    return max(1, int((e - s).total_seconds()) // 60 // base_min)


def _add_timetable(
    doc: Document,
    assignments: list[Assignment],
    perspective: str,
    entity: str,
    data: SchoolData,
) -> None:
    """Add a styled timetable table to the document for one entity."""
    days = [d.name for d in data.timeslot_config.days]
    base_min = data.timeslot_config.base_unit_minutes
    subject_colors = {s.name: s.color for s in data.subjects}
    day_to_col = {day: i + 1 for i, day in enumerate(days)}

    if perspective == "class":
        filtered = [a for a in assignments if a.school_class == entity]
    elif perspective == "teacher":
        filtered = [a for a in assignments if a.teacher == entity]
    else:
        filtered = [a for a in assignments if a.room == entity]

    lookup: dict[tuple[str, str], Assignment] = {
        (a.day, a.start_time): a for a in filtered
    }

    # Build ordered grid with None separators between sessions
    # Use first day's sessions as reference
    grid: list[tuple[str, str] | None] = []
    first_day_sessions = data.timeslot_config.days[0].sessions if data.timeslot_config.days else []
    for session in first_day_sessions:
        for s, e in _time_slots(session.start_time, session.end_time, base_min):
            grid.append((s, e))
        grid.append(None)
    while grid and grid[-1] is None:
        grid.pop()

    n_rows = 1 + len(grid)
    n_cols = 1 + len(days)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False

    # Landscape A4 with 2cm margins: ~25.7cm usable
    usable_mm = 257.0
    time_col_mm = 18.0
    day_col_mm = (usable_mm - time_col_mm) / len(days)

    for row in table.rows:
        row.cells[0].width = Mm(time_col_mm)
        for ci in range(1, n_cols):
            row.cells[ci].width = Mm(day_col_mm)

    # ── Header row — accent background, white text ──
    hdr = table.rows[0]
    _write_cell(hdr.cells[0], ["Heure"], bold=True, font_size=9, font_color=_WHITE)
    _shade_cell(hdr.cells[0], _ACCENT)
    for ci, day in enumerate(days, 1):
        _write_cell(hdr.cells[ci], [day.capitalize()], bold=True, font_size=9, font_color=_WHITE)
        _shade_cell(hdr.cells[ci], _ACCENT)

    # ── Content rows ──
    row_map: dict[str, int] = {}
    sep_set: set[int] = set()
    ri = 1
    data_row_idx = 0  # tracks alternating rows (excluding seps)
    for slot in grid:
        if slot is None:
            sep_set.add(ri)
            for ci in range(n_cols):
                _shade_cell(table.rows[ri].cells[ci], _SEP_ROW)
        else:
            start, end = slot
            row_map[start] = ri
            row_bg = _ROW_ALT if data_row_idx % 2 == 1 else _WHITE
            for ci in range(n_cols):
                _shade_cell(table.rows[ri].cells[ci], row_bg)
            _write_cell(
                table.rows[ri].cells[0],
                [f"{start}–{end}"],
                font_size=8,
                font_color="6B7280",
            )
            data_row_idx += 1
        ri += 1

    # ── Fill assignments ──
    for (day, start_time), a in lookup.items():
        if start_time not in row_map or day not in day_to_col:
            continue
        ri_a = row_map[start_time]
        ci_a = day_to_col[day]
        span = _slot_span(a.start_time, a.end_time, base_min)

        raw_hex = subject_colors.get(a.subject, "#E5E7EB")
        # Lighten subject color to 20% opacity over white for readability
        # We approximate: blend hex with white at 20% opacity
        hex_c = raw_hex.strip("#")
        try:
            r_s = int(hex_c[0:2], 16)
            g_s = int(hex_c[2:4], 16)
            b_s = int(hex_c[4:6], 16)
            # 20% subject color + 80% white
            r_l = int(r_s * 0.25 + 255 * 0.75)
            g_l = int(g_s * 0.25 + 255 * 0.75)
            b_l = int(b_s * 0.25 + 255 * 0.75)
            cell_bg = f"{r_l:02X}{g_l:02X}{b_l:02X}"
            text_color = hex_c  # full subject color for text
        except ValueError:
            cell_bg = "E5E7EB"
            text_color = _TEXT_DARK

        if perspective == "class":
            lines = [a.subject, a.teacher] + ([a.room] if a.room else [])
        elif perspective == "teacher":
            lines = [a.subject, a.school_class] + ([a.room] if a.room else [])
        else:
            lines = [a.subject, a.teacher, a.school_class]

        # Compute end row for span, skipping separator rows
        end_ri = ri_a
        steps = 0
        r = ri_a + 1
        while steps < span - 1 and r < n_rows:
            if r not in sep_set:
                steps += 1
                end_ri = r
            r += 1

        if end_ri > ri_a:
            merged = table.cell(ri_a, ci_a).merge(table.cell(end_ri, ci_a))
            _write_cell(merged, lines, bold=False, font_size=8, font_color=text_color)
            _shade_cell(merged, cell_bg)
        else:
            cell = table.rows[ri_a].cells[ci_a]
            _write_cell(cell, lines, font_size=8, font_color=text_color)
            _shade_cell(cell, cell_bg)


# ── Public API ─────────────────────────────────────────────────────────────────

def export_word(
    result: TimetableResult,
    data: SchoolData,
    output_path: str,
    perspectives: list[str] | None = None,
    logo_path: str | None = None,
) -> None:
    """Export a solved timetable to a premium landscape A4 Word document.

    Args:
        result:       Solved TimetableResult from the engine.
        data:         Original SchoolData.
        output_path:  Destination .docx path.
        perspectives: Views to include. Defaults to ['class', 'teacher'].
        logo_path:    Optional path to a school logo image (PNG/JPG).
    """
    if perspectives is None:
        perspectives = ["class", "teacher"]

    gen_date = datetime.now().strftime("%d/%m/%Y")
    doc = Document()

    # Remove default empty paragraph
    for para in list(doc.paragraphs):
        doc.element.body.remove(para._p)

    # ── Cover page (portrait, section 0) ──
    _add_cover_page(doc, data, gen_date, logo_path)

    # ── Content section (landscape) ──
    # Add a new section for landscape content
    new_section = doc.add_section(WD_ORIENT.LANDSCAPE)
    new_section.page_height = Mm(210)
    new_section.page_width = Mm(297)
    new_section.left_margin = Mm(20)
    new_section.right_margin = Mm(20)
    new_section.top_margin = Mm(18)
    new_section.bottom_margin = Mm(15)
    new_section.orientation = WD_ORIENT.LANDSCAPE

    doc_title = "Emploi du temps"
    _set_header_footer(new_section, data.school.name, doc_title, gen_date)

    first = True
    for perspective in perspectives:
        if perspective == "class":
            entities: list[tuple[str, str]] = [
                (c.name, f"Classe : {c.name}") for c in data.classes
            ]
        elif perspective == "teacher":
            entities = [(t.name, f"Enseignant : {t.name}") for t in data.teachers]
        else:
            continue

        for entity, label in entities:
            if not first:
                doc.add_page_break()
            first = False

            # Section heading
            heading_para = doc.add_paragraph()
            heading_para.paragraph_format.space_before = Pt(0)
            heading_para.paragraph_format.space_after = Pt(6)
            h_run = heading_para.add_run(label)
            h_run.font.name = _FONT_HEAD
            h_run.font.size = Pt(13)
            h_run.font.bold = True
            h_run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

            _add_timetable(doc, result.assignments, perspective, entity, data)

    doc.save(output_path)
    logger.info("Export Word premium : %s", output_path)
